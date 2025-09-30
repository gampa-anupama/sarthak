# #!/usr/bin/env python3
# import os, sys, json, uuid
# from docx import Document
# from PIL import Image
# import cv2
# import pytesseract
# import numpy as np
# from sentence_transformers import SentenceTransformer
# import faiss
# import pytesseract
# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import JSONResponse


# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # ----------------------
# # DOCX extraction
# # ----------------------
# def extract_docx(path, out_dir="output/docx_images"):
#     os.makedirs(out_dir, exist_ok=True)
#     doc = Document(path)

#     paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

#     tables = []
#     for table in doc.tables:
#         trows = []
#         for row in table.rows:
#             trows.append([cell.text.strip() for cell in row.cells])
#         tables.append(trows)

#     headers, footers = [], []
#     for i,section in enumerate(doc.sections):
#         hdr = '\n'.join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
#         ftr = '\n'.join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
#         headers.append({"section": i, "text": hdr})
#         footers.append({"section": i, "text": ftr})

#     # Extract embedded images
#     images = []
#     rels = doc.part._rels
#     for rel in rels:
#         rel_obj = rels[rel]
#         if "image" in rel_obj.target_ref:
#             blob = rel_obj.target_part.blob
#             content_type = rel_obj.target_part.content_type
#             ext = content_type.split("/")[-1]
#             filename = f"{uuid.uuid4().hex}.{ext}"
#             out_path = os.path.join(out_dir, filename)
#             with open(out_path, "wb") as f:
#                 f.write(blob)
#             images.append(out_path)

#     full_text = "\n\n".join(paragraphs)
#     meta = {
#         "source": os.path.basename(path),
#         "paragraph_count": len(paragraphs),
#         "tables_count": len(tables),
#         "headers": headers,
#         "footers": footers,
#         "images": images
#     }

#     return {"text": full_text, "paragraphs": paragraphs, "tables": tables, "meta": meta, "images": images}

# # ----------------------
# # Image OCR
# # ----------------------
# def preprocess_cv_image(path):
#     img = cv2.imread(path, cv2.IMREAD_COLOR)
#     if img is None:
#         raise FileNotFoundError(path)
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
#     gray = cv2.resize(gray, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)
#     gray = cv2.bilateralFilter(gray, 9, 75, 75)
#     th = cv2.adaptiveThreshold(gray,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
#                                cv2.THRESH_BINARY, 11, 2)
#     return th

# def ocr_with_bboxes(path, lang="eng"):
#     proc = preprocess_cv_image(path)
#     pil = Image.fromarray(proc)
#     custom_config = r'--oem 1 --psm 3'
#     data = pytesseract.image_to_data(pil, lang=lang, config=custom_config, output_type=pytesseract.Output.DICT)

#     words = []
#     full_text_lines = []
#     current_line = []

#     for i, w in enumerate(data['text']):
#         if w.strip():
#             bbox = {
#                 "text": w,
#                 "left": int(data['left'][i]),
#                 "top": int(data['top'][i]),
#                 "width": int(data['width'][i]),
#                 "height": int(data['height'][i]),
#                 "conf": int(data['conf'][i])
#             }
#             words.append(bbox)
#             current_line.append(w)
#         if i < len(data['text'])-1:
#             if data['line_num'][i] != data['line_num'][i+1]:
#                 full_text_lines.append(" ".join(current_line))
#                 current_line = []
#     if current_line:
#         full_text_lines.append(" ".join(current_line))
#     full_text = "\n".join(full_text_lines)
#     return full_text, words

# # ----------------------
# # Chunk text
# # ----------------------
# def chunk_text(text, max_chars=1000, overlap=200):
#     chunks = []
#     start = 0
#     L = len(text)
#     while start < L:
#         end = min(L, start + max_chars)
#         chunk = text[start:end].strip()
#         chunks.append({"text": chunk, "start": start, "end": end})
#         start = max(end - overlap, end)
#     return chunks

# # ----------------------
# # Main ingestion
# # ----------------------
# def main(data_folder="data", out_folder="output"):
#     os.makedirs(out_folder, exist_ok=True)
#     os.makedirs(os.path.join(out_folder, "docx_images"), exist_ok=True)
#     os.makedirs(os.path.join(out_folder, "ocr"), exist_ok=True)

#     all_chunks = []

#     model = SentenceTransformer('all-MiniLM-L6-v2')

#     # Loop through files
#     for fname in os.listdir(data_folder):
#         fpath = os.path.join(data_folder, fname)
#         ext = fname.lower().split('.')[-1]

#         if ext == "docx":
#             print(f"[DOCX] Processing {fname}...")
#             doc_out = extract_docx(fpath, out_dir=os.path.join(out_folder, "docx_images"))
#             text_chunks = chunk_text(doc_out["text"])
#             for c in text_chunks:
#                 c["source"] = fname
#                 c["type"] = "docx"
#             all_chunks.extend(text_chunks)

#             # Save text & meta
#             base = os.path.splitext(fname)[0]
#             with open(os.path.join(out_folder, f"{base}_text.txt"), "w", encoding="utf-8") as f:
#                 f.write(doc_out["text"])
#             with open(os.path.join(out_folder, f"{base}_meta.json"), "w", encoding="utf-8") as f:
#                 json.dump(doc_out["meta"], f, indent=2)

#             # OCR DOCX images
#             for img_path in doc_out["images"]:
#                 print(f"  [OCR] Extracted image {img_path}")
#                 try:
#                     ocr_text, ocr_words = ocr_with_bboxes(img_path)
#                     ocr_chunks = chunk_text(ocr_text)
#                     for c in ocr_chunks:
#                         c["source"] = os.path.basename(img_path)
#                         c["type"] = "ocr"
#                     all_chunks.extend(ocr_chunks)

#                     base_img = os.path.splitext(os.path.basename(img_path))[0]
#                     with open(os.path.join(out_folder, "ocr", f"{base_img}.txt"), "w", encoding="utf-8") as f:
#                         f.write(ocr_text)
#                     with open(os.path.join(out_folder, "ocr", f"{base_img}_bboxes.json"), "w", encoding="utf-8") as f:
#                         json.dump(ocr_words, f, indent=2)
#                 except Exception as e:
#                     print("   OCR failed:", e)

#         elif ext in ["png", "jpg", "jpeg"]:
#             print(f"[IMAGE] Processing {fname}...")
#             try:
#                 ocr_text, ocr_words = ocr_with_bboxes(fpath)
#                 ocr_chunks = chunk_text(ocr_text)
#                 for c in ocr_chunks:
#                     c["source"] = fname
#                     c["type"] = "ocr"
#                 all_chunks.extend(ocr_chunks)

#                 base_img = os.path.splitext(fname)[0]
#                 with open(os.path.join(out_folder, "ocr", f"{base_img}.txt"), "w", encoding="utf-8") as f:
#                     f.write(ocr_text)
#                 with open(os.path.join(out_folder, "ocr", f"{base_img}_bboxes.json"), "w", encoding="utf-8") as f:
#                     json.dump(ocr_words, f, indent=2)
#             except Exception as e:
#                 print("   OCR failed:", e)

#     # ----------------------
#     # Generate embeddings & build FAISS index
#     # ----------------------
#     if all_chunks:
#         print("[EMBEDDING] Generating embeddings...")
#         texts = [c["text"] for c in all_chunks]
#         embeddings = model.encode(texts, convert_to_numpy=True)

#         dim = embeddings.shape[1]
#         index = faiss.IndexFlatL2(dim)
#         index.add(embeddings)
#         faiss.write_index(index, os.path.join(out_folder, "faiss_index.idx"))

#         # Save chunk metadata
#         with open(os.path.join(out_folder, "chunks_metadata.json"), "w", encoding="utf-8") as f:
#             json.dump(all_chunks, f, indent=2)

#         print(f"[DONE] Processed {len(all_chunks)} chunks. FAISS index saved.")

# if __name__ == "__main__":
#     folder = sys.argv[1] if len(sys.argv) > 1 else "data"
#     main(folder)

#!/usr/bin/env python3
import os
import uuid
import json
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from docx import Document
from PIL import Image
import cv2
import pytesseract
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import ollama  # pip install ollama

OLLAMA_MODEL = "phi3"  # or "mistral"
TOP_K = 3
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# -------------------------
# Configuration
# -------------------------
OUT_FOLDER = "output"
DOCX_IMG_FOLDER = os.path.join(OUT_FOLDER, "docx_images")
OCR_FOLDER = os.path.join(OUT_FOLDER, "ocr")
os.makedirs(DOCX_IMG_FOLDER, exist_ok=True)
os.makedirs(OCR_FOLDER, exist_ok=True)

EMBED_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
FAISS_INDEX_PATH = os.path.join(OUT_FOLDER, "faiss_index.idx")
META_PATH = os.path.join(OUT_FOLDER, "chunks_metadata.json")

# Load existing FAISS index and metadata if they exist
if os.path.exists(FAISS_INDEX_PATH) and os.path.exists(META_PATH):
    index = faiss.read_index(FAISS_INDEX_PATH)
    with open(META_PATH, "r", encoding="utf-8") as f:
        chunks_metadata = json.load(f)
else:
    index = None
    chunks_metadata = []

# -------------------------
# FastAPI setup
# -------------------------
app = FastAPI(title="File Ingestion for RAG")

# -------------------------
# Helper functions
# -------------------------
def extract_docx(path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract embedded images
    images = []
    rels = doc.part._rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            blob = rel_obj.target_part.blob
            ext = rel_obj.target_part.content_type.split("/")[-1]
            filename = f"{uuid.uuid4().hex}.{ext}"
            out_path = os.path.join(DOCX_IMG_FOLDER, filename)
            with open(out_path, "wb") as f:
                f.write(blob)
            images.append(out_path)

    full_text = "\n\n".join(paragraphs)
    meta = {
        "source": os.path.basename(path),
        "paragraph_count": len(paragraphs),
        "images": images
    }
    return {"text": full_text, "paragraphs": paragraphs, "meta": meta, "images": images}

def preprocess_cv_image(path):
    img = cv2.imread(path, cv2.IMREAD_COLOR)
    if img is None:
        raise FileNotFoundError(path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)
    gray = cv2.bilateralFilter(gray, 9, 75, 75)
    th = cv2.adaptiveThreshold(gray,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    return th

def ocr_with_bboxes(path):
    proc = preprocess_cv_image(path)
    pil = Image.fromarray(proc)
    data = pytesseract.image_to_data(pil, lang="eng", config=r'--oem 1 --psm 3', output_type=pytesseract.Output.DICT)

    words = []
    full_text_lines = []
    current_line = []

    for i, w in enumerate(data['text']):
        if w.strip():
            bbox = {
                "text": w,
                "left": int(data['left'][i]),
                "top": int(data['top'][i]),
                "width": int(data['width'][i]),
                "height": int(data['height'][i]),
                "conf": int(data['conf'][i])
            }
            words.append(bbox)
            current_line.append(w)
        if i < len(data['text'])-1 and data['line_num'][i] != data['line_num'][i+1]:
            full_text_lines.append(" ".join(current_line))
            current_line = []
    if current_line:
        full_text_lines.append(" ".join(current_line))
    full_text = "\n".join(full_text_lines)
    return full_text, words

def chunk_text(text, max_chars=1000, overlap=200):
    chunks = []
    start = 0
    L = len(text)
    while start < L:
        end = min(L, start + max_chars)
        chunk = text[start:end].strip()
        chunks.append({"text": chunk, "start": start, "end": end})
        start = max(end - overlap, end)
    return chunks

def process_file(file_path, file_type):
    """Process uploaded file and return chunks"""
    all_chunks = []

    if file_type == "docx":
        doc_out = extract_docx(file_path)
        text_chunks = chunk_text(doc_out["text"])
        for c in text_chunks:
            c["source"] = os.path.basename(file_path)
            c["type"] = "docx"
        all_chunks.extend(text_chunks)

        # OCR images in DOCX
        for img_path in doc_out["images"]:
            try:
                ocr_text, ocr_words = ocr_with_bboxes(img_path)
                ocr_chunks = chunk_text(ocr_text)
                for c in ocr_chunks:
                    c["source"] = os.path.basename(img_path)
                    c["type"] = "ocr"
                all_chunks.extend(ocr_chunks)
            except Exception as e:
                print("OCR failed:", e)

    elif file_type in ["png", "jpg", "jpeg"]:
        try:
            ocr_text, ocr_words = ocr_with_bboxes(file_path)
            ocr_chunks = chunk_text(ocr_text)
            for c in ocr_chunks:
                c["source"] = os.path.basename(file_path)
                c["type"] = "ocr"
            all_chunks.extend(ocr_chunks)
        except Exception as e:
            print("OCR failed:", e)

    return all_chunks

def search_chunks(query: str, k: int = TOP_K):
    if not index or len(chunks_metadata) == 0:
        return []
    q_emb = EMBED_MODEL.encode([query], convert_to_numpy=True)
    distances, indices = index.search(q_emb, k)
    retrieved = [chunks_metadata[idx].get("text", "") for idx in indices[0]]
    return retrieved

# -------------------------
# Ask Ollama
# -------------------------
def ask_ollama(query: str, context: str, model: str = OLLAMA_MODEL):
    prompt = f"""
Answer the question based only on the context below.

Context:
{context}

Question:
{query}

Answer concisely:"""

    response = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
    return response["message"]["content"]

# -------------------------
# Query endpoint
# -------------------------
from fastapi import Form

@app.post("/ask")
async def ask_question(question: str = Form(...)):
    top_chunks = search_chunks(question, k=TOP_K)
    if not top_chunks:
        return JSONResponse({"error": "No chunks available. Please upload files first."}, status_code=400)
    
    context = "\n\n".join(top_chunks)
    answer = ask_ollama(question, context)
    
    return JSONResponse({
        "question": question,
        "retrieved_chunks": top_chunks,
        "answer": answer
    })


# -------------------------
# FastAPI ingestion endpoint
# -------------------------
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    # Save uploaded file
    temp_path = os.path.join("temp", f"{uuid.uuid4().hex}_{file.filename}")
    os.makedirs("temp", exist_ok=True)
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    ext = file.filename.lower().split('.')[-1]
    chunks = process_file(temp_path, ext)
    if not chunks:
        return JSONResponse({"error": "No text extracted"}, status_code=400)

    # Embeddings + FAISS
    texts = [c["text"] for c in chunks]
    embeddings = EMBED_MODEL.encode(texts, convert_to_numpy=True)
    global index, chunks_metadata
    if index is None:
        dim = embeddings.shape[1]
        index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    chunks_metadata.extend(chunks)

    # Save to disk
    os.makedirs(OUT_FOLDER, exist_ok=True)
    faiss.write_index(index, FAISS_INDEX_PATH)
    with open(META_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks_metadata, f, indent=2)

    return JSONResponse({"message": "File ingested successfully", "chunks_added": len(chunks)})
