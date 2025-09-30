#!/usr/bin/env python3
"""
ingest_docx_with_ocr.py
Usage:
    python scripts/ingest_docx_with_ocr.py data/WEEK_2_UPLOAD.docx
Output:
    output/<base>_text.txt           # combined docx text + OCR text (cleaned)
    output/<base>_meta.json          # metadata: images list + OCR bboxes/conf
    output/docx_images/<image files> # extracted embedded images
    output/ocr/<image>_ocr.txt       # OCR text per image
    output/ocr/<image>_bboxes.json   # word bbox + conf per image
"""
import os, sys, json, uuid
from docx import Document
from PIL import Image
import pytesseract
import io
import numpy as np

# If Windows and Tesseract not on PATH, set the location here:
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Optional: use OpenCV preprocessing if installed
try:
    import cv2
    HAVE_CV2 = True
except Exception:
    HAVE_CV2 = False

def extract_docx_images_and_text(docx_path, out_img_dir="output/docx_images"):
    os.makedirs(out_img_dir, exist_ok=True)
    doc = Document(docx_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # tables
    tables = []
    for table in doc.tables:
        trows = []
        for row in table.rows:
            trows.append([cell.text.strip() for cell in row.cells])
        tables.append(trows)

    # headers/footers per section
    headers, footers = [], []
    for i, section in enumerate(doc.sections):
        hdr = '\n'.join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
        ftr = '\n'.join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
        headers.append({"section": i, "text": hdr})
        footers.append({"section": i, "text": ftr})

    # extract embedded images
    images_paths = []
    rels = doc.part._rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            blob = rel_obj.target_part.blob
            content_type = rel_obj.target_part.content_type
            ext = content_type.split("/")[-1]
            filename = f"{uuid.uuid4().hex}.{ext}"
            out_path = os.path.join(out_img_dir, filename)
            with open(out_path, "wb") as f:
                f.write(blob)
            images_paths.append(out_path)

    return {
        "text": "\n\n".join(paragraphs),
        "paragraphs": paragraphs,
        "tables": tables,
        "headers": headers,
        "footers": footers,
        "images": images_paths
    }

def preprocess_image_for_ocr_cv(path):
    """OpenCV preprocessing: grayscale, upscale, denoise, threshold."""
    img = cv2.imread(path)
    if img is None:
        raise FileNotFoundError(path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # upscale small text
    gray = cv2.resize(gray, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_LINEAR)
    # denoise
    gray = cv2.bilateralFilter(gray, 9, 75, 75)
    # adaptive threshold
    th = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                               cv2.THRESH_BINARY, 11, 2)
    return th

def preprocess_image_for_ocr_pil(path):
    """Simple PIL preprocessing fallback if cv2 not available."""
    img = Image.open(path).convert("L")  # grayscale
    w, h = img.size
    # upscale moderately for small text
    img = img.resize((min(3000, int(w*2)), min(3000, int(h*2))), Image.BICUBIC)
    return img

def ocr_image(path, lang="eng", use_cv2=HAVE_CV2, psm=3):
    """
    Returns (full_text, words_list)
    words_list is list of {text,left,top,width,height,conf}
    """
    # preprocess
    if use_cv2 and HAVE_CV2:
        proc = preprocess_image_for_ocr_cv(path)
        pil_img = Image.fromarray(proc)
    else:
        pil_img = preprocess_image_for_ocr_pil(path)

    config = f'--oem 1 --psm {psm}'  # OEM 1 = LSTM engine; change psm for layouts
    data = pytesseract.image_to_data(pil_img, lang=lang, config=config, output_type=pytesseract.Output.DICT)
    # reconstruct lines
    lines = []
    current = []
    words = []
    n = len(data['text'])
    for i in range(n):
        txt = str(data['text'][i]).strip()
        if txt:
            words.append({
                "text": txt,
                "left": int(data['left'][i]),
                "top": int(data['top'][i]),
                "width": int(data['width'][i]),
                "height": int(data['height'][i]),
                "conf": int(float(data['conf'][i])) if data['conf'][i] != '-1' else -1
            })
            current.append(txt)
        # detect line break
        if i < n-1 and data['line_num'][i] != data['line_num'][i+1]:
            if current:
                lines.append(" ".join(current))
            current = []
    if current:
        lines.append(" ".join(current))
    full_text = "\n".join(lines).strip()
    return full_text, words

def clean_text(s):
    # simple normalizations: unify whitespace, remove rare control chars
    if not s:
        return ""
    s = s.replace('\r', ' ')
    s = s.replace('\t', ' ')
    s = s.strip()
    # collapse multiple spaces/newlines
    while '  ' in s:
        s = s.replace('  ', ' ')
    return s

def main():
    if len(sys.argv) < 2:
        print("Usage: python scripts/ingest_docx_with_ocr.py path/to/file.docx")
        sys.exit(1)
    path = sys.argv[1]
    base = os.path.splitext(os.path.basename(path))[0]
    os.makedirs("output", exist_ok=True)
    os.makedirs("output/docx_images", exist_ok=True)
    os.makedirs("output/ocr", exist_ok=True)

    # 1) Extract docx text + embedded images
    docinfo = extract_docx_images_and_text(path, out_img_dir="output/docx_images")
    doc_text = docinfo["text"] or ""

    # 2) For each extracted image -> OCR
    images_meta = []
    for img_path in docinfo["images"]:
        try:
            ocr_text, words = ocr_image(img_path)
        except Exception as e:
            print(f"WARNING: OCR failed for {img_path}: {e}")
            ocr_text, words = "", []
        # save individual OCR files
        img_name = os.path.splitext(os.path.basename(img_path))[0]
        txt_path = f"output/ocr/{img_name}_ocr.txt"
        json_path = f"output/ocr/{img_name}_bboxes.json"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(ocr_text)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(words, f, indent=2)
        images_meta.append({
            "image_path": img_path,
            "ocr_text_file": txt_path,
            "ocr_bboxes_file": json_path,
            "ocr_text": ocr_text,
            "word_count": len(words)
        })

    # 3) Combine doc text + image OCR (append with markers) and clean
    combined_parts = []
    if doc_text.strip():
        combined_parts.append("=== DOCX EXTRACTED TEXT ===\n" + doc_text.strip())
    for im in images_meta:
        if im["ocr_text"].strip():
            combined_parts.append(f"=== OCR from image: {os.path.basename(im['image_path'])} ===\n" + im["ocr_text"].strip())
    combined_text = "\n\n".join(combined_parts)
    combined_text = clean_text(combined_text)

    # 4) Save outputs
    out_text_path = f"output/{base}_text.txt"
    out_meta_path = f"output/{base}_meta.json"
    with open(out_text_path, "w", encoding="utf-8") as f:
        f.write(combined_text)
    meta = {
        "source": os.path.basename(path),
        "paragraph_count": len(docinfo.get("paragraphs", [])),
        "tables_count": len(docinfo.get("tables", [])),
        "images": images_meta,
        "headers": docinfo.get("headers"),
        "footers": docinfo.get("footers")
    }
    with open(out_meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2)

    print("Done.")
    print("Combined text saved to:", out_text_path)
    print("Metadata saved to:", out_meta_path)
    if images_meta:
        print("Image OCR outputs in output/ocr/")

if __name__ == "__main__":
    main()
