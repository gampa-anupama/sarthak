# # # #!/usr/bin/env python3
# # # import os
# # # import shutil
# # # from fastapi import FastAPI, UploadFile, File
# # # from fastapi.responses import JSONResponse
# # # import uvicorn

# # # # Import your existing ingestion code
# # # import ingest_single

# # # app = FastAPI()

# # # UPLOAD_FOLDER = "data"
# # # os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # # @app.post("/upload/")
# # # async def upload_file(file: UploadFile = File(...)):
# # #     """
# # #     Upload a DOCX or Image file.
# # #     """
# # #     filename = os.path.join(UPLOAD_FOLDER, file.filename)
# # #     with open(filename, "wb") as buffer:
# # #         shutil.copyfileobj(file.file, buffer)

# # #     # Run your ingestion pipeline
# # #     try:
# # #         ingest_single.main(UPLOAD_FOLDER, out_folder="output")
# # #         return JSONResponse(content={
# # #             "status": "success",
# # #             "message": f"File {file.filename} uploaded and processed successfully."
# # #         })
# # #     except Exception as e:
# # #         return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)

# # # if __name__ == "__main__":
# # #     uvicorn.run(app, host="0.0.0.0", port=8000)

# # #!/usr/bin/env python3
# # import os, json, uuid
# # from fastapi import FastAPI, UploadFile, File
# # from fastapi.responses import JSONResponse
# # from docx import Document
# # from PIL import Image
# # import cv2
# # import pytesseract
# # import numpy as np
# # from sentence_transformers import SentenceTransformer
# # import faiss
# # import uvicorn

# # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # app = FastAPI()

# # # ----------------------
# # # DOCX extraction
# # # ----------------------
# # def extract_docx(path, out_dir="output/docx_images"):
# #     os.makedirs(out_dir, exist_ok=True)
# #     doc = Document(path)
# #     paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# #     # Extract embedded images
# #     images = []
# #     rels = doc.part._rels
# #     for rel in rels:
# #         rel_obj = rels[rel]
# #         if "image" in rel_obj.target_ref:
# #             blob = rel_obj.target_part.blob
# #             content_type = rel_obj.target_part.content_type
# #             ext = content_type.split("/")[-1]
# #             filename = f"{uuid.uuid4().hex}.{ext}"
# #             out_path = os.path.join(out_dir, filename)
# #             with open(out_path, "wb") as f:
# #                 f.write(blob)
# #             images.append(out_path)

# #     full_text = "\n\n".join(paragraphs)
# #     return {"text": full_text, "paragraphs": paragraphs, "images": images}

# # # ----------------------
# # # Image OCR
# # # ----------------------
# # def preprocess_cv_image(path):
# #     img = cv2.imread(path, cv2.IMREAD_COLOR)
# #     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
# #     gray = cv2.resize(gray, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)
# #     gray = cv2.bilateralFilter(gray, 9, 75, 75)
# #     th = cv2.adaptiveThreshold(gray,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
# #                                cv2.THRESH_BINARY, 11, 2)
# #     return th

# # def ocr_with_bboxes(path, lang="eng"):
# #     proc = preprocess_cv_image(path)
# #     pil = Image.fromarray(proc)
# #     custom_config = r'--oem 1 --psm 3'
# #     data = pytesseract.image_to_string(pil, lang=lang, config=custom_config)
# #     return data.strip()

# # # ----------------------
# # # Chunk text
# # # ----------------------
# # def chunk_text(text, max_chars=1000, overlap=200):
# #     chunks = []
# #     start = 0
# #     L = len(text)
# #     while start < L:
# #         end = min(L, start + max_chars)
# #         chunk = text[start:end].strip()
# #         if chunk:
# #             chunks.append(chunk)
# #         start = max(end - overlap, end)
# #     return chunks

# # # ----------------------
# # # Upload API
# # # ----------------------
# # @app.post("/upload/")
# # async def upload_file(file: UploadFile = File(...)):
# #     out_folder = "output"
# #     os.makedirs(out_folder, exist_ok=True)
# #     os.makedirs(os.path.join(out_folder, "ocr"), exist_ok=True)
# #     os.makedirs(os.path.join(out_folder, "docx_images"), exist_ok=True)

# #     file_path = os.path.join(out_folder, file.filename)
# #     with open(file_path, "wb") as f:
# #         f.write(await file.read())

# #     all_chunks = []
# #     model = SentenceTransformer('all-MiniLM-L6-v2')

# #     ext = file.filename.lower().split('.')[-1]
# #     if ext == "docx":
# #         doc_out = extract_docx(file_path, out_dir=os.path.join(out_folder, "docx_images"))

# #         # ✅ PRINT extracted DOCX text
# #         print(f"\n[DOCX] Extracted Text from {file.filename}:\n{'='*60}\n{doc_out['text']}\n{'='*60}\n")

# #         chunks = chunk_text(doc_out["text"])
# #         all_chunks.extend(chunks)

# #         # Process images inside DOCX
# #         for img_path in doc_out["images"]:
# #             try:
# #                 ocr_text = ocr_with_bboxes(img_path)
# #                 print(f"\n[OCR] Extracted from {img_path}:\n{'-'*60}\n{ocr_text}\n{'-'*60}\n")
# #                 chunks = chunk_text(ocr_text)
# #                 all_chunks.extend(chunks)
# #             except Exception as e:
# #                 print("OCR failed:", e)

# #     elif ext in ["png", "jpg", "jpeg"]:
# #         ocr_text = ocr_with_bboxes(file_path)

# #         # ✅ PRINT extracted OCR text from image
# #         print(f"\n[OCR] Extracted from {file.filename}:\n{'-'*60}\n{ocr_text}\n{'-'*60}\n")

# #         chunks = chunk_text(ocr_text)
# #         all_chunks.extend(chunks)

# #     # Build FAISS index
# #     if all_chunks:
# #         texts = all_chunks
# #         embeddings = model.encode(texts, convert_to_numpy=True)
# #         dim = embeddings.shape[1]
# #         index = faiss.IndexFlatL2(dim)
# #         index.add(embeddings)
# #         faiss.write_index(index, os.path.join(out_folder, "faiss_index.idx"))
# #         with open(os.path.join(out_folder, "chunks_metadata.json"), "w", encoding="utf-8") as f:
# #             json.dump(all_chunks, f, indent=2)

# #     return JSONResponse(content={"message": "File processed", "chunks_count": len(all_chunks)})

# # if __name__ == "__main__":
# #     uvicorn.run(app, host="0.0.0.0", port=8000)


# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import JSONResponse
# import os, json, uuid, shutil
# from scripts.ingest_single import main, extract_docx, ocr_with_bboxes, chunk_text  # import from your script

# app = FastAPI()

# UPLOAD_DIR = "uploads"
# OUTPUT_DIR = "output"
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)


# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     """
#     Upload a DOCX or Image -> extract text, OCR if needed, return extracted content
#     """
#     # Save uploaded file
#     file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex}_{file.filename}")
#     with open(file_path, "wb") as buffer:
#         shutil.copyfileobj(file.file, buffer)

#     ext = file.filename.lower().split(".")[-1]

#     try:
#         if ext == "docx":
#             # Extract from DOCX
#             doc_out = extract_docx(file_path, out_dir=os.path.join(OUTPUT_DIR, "docx_images"))
#             return JSONResponse({
#                 "filename": file.filename,
#                 "type": "docx",
#                 "text": doc_out["text"],
#                 "paragraphs": doc_out["paragraphs"],
#                 "tables": doc_out["tables"],
#                 "meta": doc_out["meta"]
#             })

#         elif ext in ["jpg", "jpeg", "png"]:
#             # OCR for image
#             ocr_text, ocr_words = ocr_with_bboxes(file_path)
#             return JSONResponse({
#                 "filename": file.filename,
#                 "type": "image",
#                 "extracted_text": ocr_text,
#                 "bboxes": ocr_words
#             })

#         else:
#             return JSONResponse({"error": "Unsupported file type"}, status_code=400)

#     except Exception as e:
#         return JSONResponse({"error": str(e)}, status_code=500)


# @app.get("/search/")
# async def search(query: str):
#     """
#     Future step: Query FAISS index (not yet integrated in this snippet).
#     """
#     return {"message": f"Searching for: {query}"}


#!/usr/bin/env python3
import os
import uuid
import json
from fastapi import FastAPI, Form, UploadFile, File
from fastapi.responses import JSONResponse
from docx import Document
from PIL import Image
import cv2
import pytesseract
from sentence_transformers import SentenceTransformer
import faiss
import ollama  # pip install ollama

# -------------------------
# Configuration
# -------------------------
OUT_FOLDER = "output1"
DOCX_IMG_FOLDER = os.path.join(OUT_FOLDER, "docx_images")
OCR_FOLDER = os.path.join(OUT_FOLDER, "ocr")
os.makedirs(DOCX_IMG_FOLDER, exist_ok=True)
os.makedirs(OCR_FOLDER, exist_ok=True)
os.makedirs("temp", exist_ok=True)

EMBED_MODEL_NAME = 'all-MiniLM-L6-v2'
OLLAMA_MODEL = 'phi3'  # or 'mistral'
TOP_K = 3

FAISS_INDEX_PATH = os.path.join(OUT_FOLDER, "faiss_index.idx")
CHUNKS_META_PATH = os.path.join(OUT_FOLDER, "chunks_metadata.json")

# -------------------------
# Initialize embedding model
# -------------------------
embed_model = SentenceTransformer(EMBED_MODEL_NAME)

# -------------------------
# Load FAISS index + metadata if available
# -------------------------
if os.path.exists(FAISS_INDEX_PATH) and os.path.exists(CHUNKS_META_PATH):
    index = faiss.read_index(FAISS_INDEX_PATH)
    with open(CHUNKS_META_PATH, "r", encoding="utf-8") as f:
        chunks_metadata = json.load(f)
    # Ensure dict
    for i, item in enumerate(chunks_metadata):
        if isinstance(item, str):
            chunks_metadata[i] = json.loads(item)
else:
    index = None
    chunks_metadata = []

# -------------------------
# FastAPI setup
# -------------------------
app = FastAPI(title="RAG QA & Ingestion API")

# -------------------------
# Helper functions
# -------------------------
def extract_docx(path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract embedded images
    images = []
    rels = doc.part._rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            blob = rel_obj.target_part.blob
            ext = rel_obj.target_part.content_type.split("/")[-1]
            filename = f"{uuid.uuid4().hex}.{ext}"
            out_path = os.path.join(DOCX_IMG_FOLDER, filename)
            with open(out_path, "wb") as f:
                f.write(blob)
            images.append(out_path)

    full_text = "\n\n".join(paragraphs)
    meta = {"source": os.path.basename(path), "paragraph_count": len(paragraphs), "images": images}
    return {"text": full_text, "paragraphs": paragraphs, "meta": meta, "images": images}

def preprocess_cv_image(path):
    img = cv2.imread(path, cv2.IMREAD_COLOR)
    if img is None:
        raise FileNotFoundError(path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)
    gray = cv2.bilateralFilter(gray, 9, 75, 75)
    th = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                               cv2.THRESH_BINARY, 11, 2)
    return th

def ocr_with_bboxes(path):
    proc = preprocess_cv_image(path)
    pil = Image.fromarray(proc)
    data = pytesseract.image_to_data(pil, lang="eng", config=r'--oem 1 --psm 3', output_type=pytesseract.Output.DICT)

    words = []
    full_text_lines = []
    current_line = []

    for i, w in enumerate(data['text']):
        if w.strip():
            bbox = {"text": w, "left": int(data['left'][i]), "top": int(data['top'][i]),
                    "width": int(data['width'][i]), "height": int(data['height'][i]),
                    "conf": int(data['conf'][i])}
            words.append(bbox)
            current_line.append(w)
        if i < len(data['text'])-1 and data['line_num'][i] != data['line_num'][i+1]:
            full_text_lines.append(" ".join(current_line))
            current_line = []
    if current_line:
        full_text_lines.append(" ".join(current_line))
    full_text = "\n".join(full_text_lines)
    return full_text, words

def chunk_text(text, max_chars=1000, overlap=200):
    chunks = []
    start = 0
    L = len(text)
    while start < L:
        end = min(L, start + max_chars)
        chunk = text[start:end].strip()
        chunks.append({"text": chunk, "start": start, "end": end})
        start = max(end - overlap, end)
    return chunks

def process_file(file_path, file_type):
    all_chunks = []
    if file_type == "docx":
        doc_out = extract_docx(file_path)
        text_chunks = chunk_text(doc_out["text"])
        for c in text_chunks:
            c["source"] = os.path.basename(file_path)
            c["type"] = "docx"
        all_chunks.extend(text_chunks)

        # OCR images
        for img_path in doc_out["images"]:
            try:
                ocr_text, _ = ocr_with_bboxes(img_path)
                ocr_chunks = chunk_text(ocr_text)
                for c in ocr_chunks:
                    c["source"] = os.path.basename(img_path)
                    c["type"] = "ocr"
                all_chunks.extend(ocr_chunks)
            except Exception as e:
                print("OCR failed:", e)

    elif file_type in ["png", "jpg", "jpeg"]:
        try:
            ocr_text, _ = ocr_with_bboxes(file_path)
            ocr_chunks = chunk_text(ocr_text)
            for c in ocr_chunks:
                c["source"] = os.path.basename(file_path)
                c["type"] = "ocr"
            all_chunks.extend(ocr_chunks)
        except Exception as e:
            print("OCR failed:", e)
    return all_chunks

# -------------------------
# Upload ingestion endpoint
# -------------------------
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    fname = file.filename
    ext = fname.lower().split('.')[-1]
    if ext not in ["docx", "png", "jpg", "jpeg"]:
        return JSONResponse({"error": "Unsupported file type"}, status_code=400)

    temp_path = os.path.join("temp", f"{uuid.uuid4().hex}_{fname}")
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    chunks = process_file(temp_path, ext)
    if not chunks:
        return JSONResponse({"error": "No text extracted from file"}, status_code=400)

    # Add embeddings to FAISS
    texts = [c["text"] for c in chunks]
    embeddings = embed_model.encode(texts, convert_to_numpy=True)
    global index, chunks_metadata
    if index is None:
        dim = embeddings.shape[1]
        index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    chunks_metadata.extend(chunks)

    # Save updated index & metadata
    faiss.write_index(index, FAISS_INDEX_PATH)
    with open(CHUNKS_META_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks_metadata, f, indent=2)

    return JSONResponse({"message": "File ingested successfully", "chunks_added": len(chunks)})

# -------------------------
# Query endpoint using Ollama
# -------------------------
def search_chunks(query: str, k: int = TOP_K):
    q_emb = embed_model.encode([query], convert_to_numpy=True)
    distances, indices = index.search(q_emb, k)
    retrieved = [chunks_metadata[idx].get("text", "") for idx in indices[0]]
    return retrieved

def ask_ollama(query: str, context: str, model: str = OLLAMA_MODEL):
    prompt = f"""
Answer the question based only on the context below.

Context:
{context}

Question:
{query}

Answer concisely:"""

    response = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
    return response["message"]["content"]

@app.post("/ask")
async def ask_question(question: str = Form(...)):
    top_chunks = search_chunks(question, k=TOP_K)
    context = "\n\n".join(top_chunks)
    answer = ask_ollama(question, context)
    return JSONResponse({"question": question, "retrieved_chunks": top_chunks, "answer": answer})
